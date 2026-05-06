"""RAG Document Knowledge Base REST API Router.

Provides document ingestion, chunking, TF-IDF semantic search, and persistent
storage for building a local RAG pipeline. Integrates with the Knowledge Graph
by auto-extracting key terms as entities.

Endpoints
---------
POST   /api/electron/kb/documents              – upload document
GET    /api/electron/kb/documents              – list documents
GET    /api/electron/kb/documents/{id}         – get document with chunks
GET    /api/electron/kb/documents/{id}/download – download original file
PATCH  /api/electron/kb/documents/{id}         – update document metadata
DELETE /api/electron/kb/documents/{id}         – delete document
POST   /api/electron/kb/documents/batch-delete – batch delete documents
POST   /api/electron/kb/search                 – semantic search
POST   /api/electron/kb/search/hybrid          – hybrid TF-IDF + embedding search
GET    /api/electron/kb/tags                   – list tags and categories
GET    /api/electron/kb/embeddings/status      – embedding provider status
GET    /api/electron/kb/stats                  – knowledge base statistics
POST   /api/electron/kb/reindex                – force re-index all documents
"""

from __future__ import annotations

import asyncio
import json
import logging
import os
import re
import time
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

import numpy as np
from fastapi import APIRouter, File, HTTPException, Query, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from deerflow.config.paths import get_paths

from app.gateway.routers.embeddings import (
    compute_similarity,
    get_embedding_provider,
    init_embedding_provider,
)

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/electron/kb", tags=["knowledge-base"])

_lock = asyncio.Lock()

# ── Constants ─────────────────────────────────────────────────────────

CHUNK_SIZE = 600       # target characters per chunk
CHUNK_OVERLAP = 100    # character overlap between adjacent chunks
SUPPORTED_EXTS = {".txt", ".md", ".py", ".js", ".ts", ".json", ".html", ".css", ".yaml", ".yml", ".xml", ".csv", ".log", ".rst"}
MAX_DOC_SIZE_MB = 20

# ── In-memory state ───────────────────────────────────────────────────

_documents: dict[str, dict] = {}       # doc_id -> metadata dict
_chunks: dict[str, list[dict]] = {}    # doc_id -> [chunk dicts]
_vectorizer: TfidfVectorizer | None = None
_chunk_vectors: np.ndarray | None = None
_chunk_index: list[dict] = []          # [{doc_id, chunk_idx, text}, ...]
_persistence_path: Optional[Path] = None
_docs_dir: Optional[Path] = None
_vectors_path: Optional[Path] = None
_chunk_vectors_np: np.ndarray | None = None  # shape (N, D) float32 embedding vectors

# ── Pydantic models ───────────────────────────────────────────────────


class DocumentMeta(BaseModel):
    id: str = Field(..., description="Document ID")
    filename: str = Field(..., description="Original filename")
    title: str = Field(default="", description="Document title")
    fileType: str = Field(default="", description="File extension")
    fileSize: int = Field(default=0, description="File size in bytes")
    chunkCount: int = Field(default=0, description="Number of chunks")
    charCount: int = Field(default=0, description="Total characters")
    pageCount: int | None = Field(default=None, description="Number of pages (PDF/DOCX only)")
    tags: list[str] = Field(default_factory=list, description="User-defined tags")
    category: str = Field(default="general", description="Document category")
    createdAt: str = Field(default="", description="ISO timestamp")
    updatedAt: str = Field(default="", description="ISO timestamp")


class DocumentDetail(DocumentMeta):
    chunks: list[dict] = Field(default_factory=list, description="Document chunks with text")


class DocumentListResponse(BaseModel):
    documents: list[DocumentMeta]
    total: int


class SearchRequest(BaseModel):
    query: str = Field(..., min_length=1, description="Search query text")
    topK: int = Field(default=10, ge=1, le=100, description="Number of results")
    docId: str | None = Field(default=None, description="Optional: limit search to one document")
    minScore: float = Field(default=0.05, ge=0.0, le=1.0, description="Minimum similarity score")


class SearchResult(BaseModel):
    docId: str = ""
    docTitle: str = ""
    docFilename: str = ""
    chunkIdx: int = 0
    chunkText: str = ""
    score: float = 0.0
    preview: str = Field(default="", description="Preview with highlighted context")


class SearchResponse(BaseModel):
    results: list[SearchResult]
    totalChunks: int = 0
    queryTimeMs: float = 0.0


class KBStatsResponse(BaseModel):
    totalDocuments: int = 0
    totalChunks: int = 0
    totalChars: int = 0
    fileTypes: dict[str, int] = Field(default_factory=dict)
    indexedVectors: int = 0
    avgChunkSize: float = 0.0
    categories: dict[str, int] = Field(default_factory=dict)
    tags: dict[str, int] = Field(default_factory=dict)


class DocumentUpdateRequest(BaseModel):
    title: str | None = Field(default=None, description="New document title")
    tags: list[str] | None = Field(default=None, description="New tags list")
    category: str | None = Field(default=None, description="New category")


class TagsResponse(BaseModel):
    tags: list[dict] = Field(default_factory=list)  # [{name, count}, ...]
    categories: list[dict] = Field(default_factory=list)  # [{name, count}, ...]


class HybridSearchRequest(BaseModel):
    query: str = Field(..., min_length=1)
    topK: int = Field(default=10, ge=1, le=100)
    docId: str | None = Field(default=None)
    alpha: float = Field(default=0.6, ge=0.0, le=1.0, description="Embedding weight (0=TF-IDF only, 1=embedding only)")
    minScore: float = Field(default=0.05, ge=0.0, le=1.0)


class HybridSearchResult(BaseModel):
    docId: str = ""
    docTitle: str = ""
    docFilename: str = ""
    chunkIdx: int = 0
    chunkText: str = ""
    score: float = 0.0       # hybrid score (final)
    tfidfScore: float = 0.0  # TF-IDF component
    embeddingScore: float = 0.0  # embedding component
    preview: str = ""


class HybridSearchResponse(BaseModel):
    results: list[HybridSearchResult]
    totalChunks: int = 0
    queryTimeMs: float = 0.0
    searchMode: str = "hybrid"  # "hybrid" | "tfidf_only" | "embedding_only"
    alpha: float = 0.6


class EmbeddingStatusResponse(BaseModel):
    provider: str = "none"
    model: str = "N/A"
    dimension: int = 0
    available: bool = False
    embeddedChunks: int = 0
    message: str = ""


class BatchDeleteRequest(BaseModel):
    ids: list[str] = Field(..., min_length=1, description="Document IDs to delete")


class BatchDeleteResponse(BaseModel):
    success: bool = True
    deleted: int = 0
    failed: int = 0
    errors: list[str] = Field(default_factory=list)


class BatchUpdateRequest(BaseModel):
    ids: list[str] = Field(..., min_length=1, description="Document IDs to update")
    tags: list[str] | None = Field(default=None, description="Tags to apply")
    category: str | None = Field(default=None, description="Category to apply")
    mode: str = Field(default="set", description="Mode: 'set', 'add', or 'remove'")


class BatchUpdateResponse(BaseModel):
    success: bool = True
    updated: int = 0
    failed: int = 0
    errors: list[str] = Field(default_factory=list)


# ── Persistence helpers ───────────────────────────────────────────────


def _get_persistence_path() -> Path:
    global _persistence_path
    if _persistence_path is None:
        try:
            base = get_paths().base_dir
        except Exception:
            base = Path(".")
        _persistence_path = base / "knowledge_base.json"
    return _persistence_path


def _get_docs_dir() -> Path:
    global _docs_dir
    if _docs_dir is None:
        try:
            base = get_paths().base_dir
        except Exception:
            base = Path(".")
        _docs_dir = base / "knowledge_docs"
    _docs_dir.mkdir(parents=True, exist_ok=True)
    return _docs_dir


def _get_vectors_path() -> Path:
    global _vectors_path
    if _vectors_path is None:
        _vectors_path = _get_persistence_path().parent / "knowledge_base_vectors.npy"
    return _vectors_path


def _load_vectors() -> np.ndarray | None:
    """Load embedding vectors from NPY binary file (fast + compact)."""
    vp = _get_vectors_path()
    if vp.exists():
        try:
            data = np.load(vp)
            if data.ndim == 2 and data.shape[0] > 0 and data.shape[1] > 0:
                logger.info("Loaded %d embedding vectors (dim=%d) from %s", data.shape[0], data.shape[1], vp.name)
                return data.astype(np.float32)
        except Exception:
            logger.exception("Failed to load embedding vectors from %s", vp)
    return None


def _save_vectors(vectors: np.ndarray) -> None:
    """Save embedding vectors to NPY binary file."""
    vp = _get_vectors_path()
    try:
        np.save(vp, vectors.astype(np.float32))
        logger.debug("Saved %d embedding vectors to %s", vectors.shape[0], vp.name)
    except Exception:
        logger.exception("Failed to save embedding vectors to %s", vp)


async def _load_state() -> None:
    """Load knowledge base from persistent JSON file on startup."""
    path = _get_persistence_path()
    if not path.exists():
        return
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        logger.exception("Failed to read knowledge_base.json")
        return

    docs_raw = data.get("documents", [])
    for d in docs_raw:
        if isinstance(d, dict) and d.get("id"):
            _documents[d["id"]] = d

    chunks_raw = data.get("chunks", {})
    for doc_id, clist in chunks_raw.items():
        if doc_id in _documents:
            _chunks[doc_id] = clist

    logger.info("Loaded %d documents and chunks for %d docs from knowledge_base.json",
                len(_documents), len(_chunks))

    # Rebuild vector index
    if _documents and _chunks:
        _rebuild_index()


async def _save_state() -> None:
    """Persist knowledge base state to JSON file."""
    path = _get_persistence_path()
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        data = {
            "documents": list(_documents.values()),
            "chunks": {doc_id: clist for doc_id, clist in _chunks.items()},
            "updatedAt": time.time(),
        }
        path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
    except Exception:
        logger.exception("Failed to persist knowledge base state")


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


# ── Text extraction ───────────────────────────────────────────────────


def _extract_text(filepath: str, ext: str) -> str | None:
    """Extract text content from a file based on its extension."""
    try:
        # Plain text files
        if ext in {".txt", ".md", ".py", ".js", ".ts", ".json", ".html",
                   ".css", ".yaml", ".yml", ".xml", ".csv", ".log", ".rst"}:
            return Path(filepath).read_text(encoding="utf-8", errors="replace")

        # PDF extraction
        if ext == ".pdf":
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(filepath)
                pages = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
                return "\n\n".join(pages) if pages else None
            except ImportError:
                logger.warning("PyPDF2 not available, cannot extract PDF text: %s", filepath)
                return f"[PDF file: {os.path.basename(filepath)} — install PyPDF2 for text extraction]"
            except Exception as exc:
                logger.warning("Failed to extract PDF text from %s: %s", filepath, exc)
                return f"[PDF extraction error: {os.path.basename(filepath)}]"

        # Word documents
        if ext == ".docx":
            try:
                from docx import Document
                doc = Document(filepath)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                return "\n\n".join(paragraphs) if paragraphs else None
            except ImportError:
                logger.warning("python-docx not available: %s", filepath)
                return f"[DOCX file: {os.path.basename(filepath)} — install python-docx for text extraction]"
            except Exception as exc:
                logger.warning("Failed to extract DOCX from %s: %s", filepath, exc)
                return f"[DOCX extraction error: {os.path.basename(filepath)}]"

        return None

    except Exception as exc:
        logger.error("Failed to extract text from %s: %s", filepath, exc)
        return None


def _get_page_count(filepath: str, ext: str) -> int | None:
    """Get page count for PDF or DOCX files."""
    try:
        if ext == ".pdf":
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(filepath)
                return len(reader.pages)
            except ImportError:
                return None
            except Exception:
                return None
        if ext == ".docx":
            try:
                from docx import Document
                doc = Document(filepath)
                return len(doc.sections)
            except ImportError:
                return None
            except Exception:
                return None
        return None
    except Exception:
        return None


# ── Smart chunking ─────────────────────────────────────────────────────


def _smart_chunk(text: str) -> list[str]:
    """Split text into overlapping chunks preserving paragraph boundaries."""
    if not text.strip():
        return []

    # Clean excessive whitespace
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)

    paragraphs = text.split("\n\n")
    chunks: list[str] = []
    current = ""
    overlap_buffer = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # If paragraph is very long, split by sentences
        if len(para) > CHUNK_SIZE:
            sentences = re.split(r"(?<=[.!?。！？])\s+", para)
            para_parts = []
            buf = ""
            for sent in sentences:
                if len(buf) + len(sent) <= CHUNK_SIZE:
                    buf = (buf + " " + sent).strip()
                else:
                    if buf:
                        para_parts.append(buf)
                    buf = sent
            if buf:
                para_parts.append(buf)

            for part in para_parts:
                if len(current) + len(part) <= CHUNK_SIZE:
                    current = (current + "\n\n" + part).strip()
                else:
                    if current:
                        chunks.append(current)
                        # Keep overlap
                        overlap_buffer = current[-CHUNK_OVERLAP:] if len(current) > CHUNK_OVERLAP else current
                    current = overlap_buffer + "\n\n" + part if overlap_buffer else part
        else:
            if len(current) + len(para) <= CHUNK_SIZE:
                current = (current + "\n\n" + para).strip()
            else:
                if current:
                    chunks.append(current)
                    overlap_buffer = current[-CHUNK_OVERLAP:] if len(current) > CHUNK_OVERLAP else current
                current = overlap_buffer + "\n\n" + para if overlap_buffer else para

    if current:
        chunks.append(current)

    return chunks


# ── TF-IDF vectorization ───────────────────────────────────────────────


def _rebuild_index() -> None:
    """Build TF-IDF vector index from all chunks in memory.

    Also loads embedding vectors from the NPY binary file so they are
    aligned with the current chunk_index ordering.
    """
    global _vectorizer, _chunk_vectors, _chunk_index, _chunk_vectors_np

    _chunk_index = []
    all_texts: list[str] = []

    for doc_id, clist in _chunks.items():
        for idx, chunk in enumerate(clist):
            text = chunk.get("text", "")
            if text.strip():
                all_texts.append(text)
                _chunk_index.append({"doc_id": doc_id, "chunk_idx": idx, "text": text})

    if not all_texts:
        _vectorizer = None
        _chunk_vectors = None
        _chunk_vectors_np = None
        return

    _vectorizer = TfidfVectorizer(
        max_features=10000,
        stop_words="english",
        ngram_range=(1, 2),
        lowercase=True,
    )
    _chunk_vectors = _vectorizer.fit_transform(all_texts)

    # Load embedding vectors (aligned with chunk_index)
    _chunk_vectors_np = _load_vectors()
    if _chunk_vectors_np is not None and _chunk_vectors_np.shape[0] != len(_chunk_index):
        logger.warning(
            "Vector count mismatch: %d vectors vs %d chunks — discarding vectors",
            _chunk_vectors_np.shape[0], len(_chunk_index),
        )
        _chunk_vectors_np = None

    logger.info("Rebuilt TF-IDF index: %d chunks, %d features%s",
                len(all_texts), _chunk_vectors.shape[1],
                f", {_chunk_vectors_np.shape[0]} embedding vectors" if _chunk_vectors_np is not None else "")


def _search_chunks(query: str, top_k: int = 10, doc_id: str | None = None,
                   min_score: float = 0.05) -> list[SearchResult]:
    """Search chunks using TF-IDF cosine similarity."""
    global _vectorizer, _chunk_vectors, _chunk_index

    if _vectorizer is None or _chunk_vectors is None or not _chunk_index:
        return []

    query_vec = _vectorizer.transform([query])
    scores = cosine_similarity(query_vec, _chunk_vectors)[0]

    results: list[SearchResult] = []
    for i, score in enumerate(scores):
        if score < min_score:
            continue
        entry = _chunk_index[i]
        if doc_id and entry["doc_id"] != doc_id:
            continue

        doc = _documents.get(entry["doc_id"], {})
        text = entry["text"]

        # Build preview with context highlighting
        preview = text[:300]
        if len(text) > 300:
            preview += "..."

        results.append(SearchResult(
            docId=entry["doc_id"],
            docTitle=doc.get("title", ""),
            docFilename=doc.get("filename", ""),
            chunkIdx=entry["chunk_idx"],
            chunkText=text,
            score=round(float(score), 4),
            preview=preview,
        ))

    # Sort by score descending
    results.sort(key=lambda x: x.score, reverse=True)
    return results[:top_k]


# ── Auto-entity extraction (lightweight) ───────────────────────────────


def _extract_key_terms(text: str, max_terms: int = 10) -> list[str]:
    """Extract significant capitalized terms and noun phrases as potential entities.

    This is a lightweight, heuristic-based extractor that feeds the Knowledge Graph.
    """
    if _vectorizer is None:
        return []

    # Get term importance from TF-IDF
    try:
        vec = _vectorizer.transform([text])
        feature_names = _vectorizer.get_feature_names_out()
        tfidf_scores = vec.toarray()[0]

        # Get top terms by TF-IDF weight
        top_indices = np.argsort(tfidf_scores)[-max_terms:][::-1]
        terms = []
        for idx in top_indices:
            if tfidf_scores[idx] > 0:
                term = feature_names[idx]
                # Keep multi-word terms and capitalized terms
                if " " in term or term[0].isupper():
                    terms.append(term)
        return terms[:max_terms]
    except Exception:
        return []


async def _link_to_knowledge_graph(doc_id: str, title: str, text: str) -> int:
    """Extract key terms from document and create entities in knowledge graph."""
    try:
        from app.gateway.routers.knowledge_graph import (
            CreateEntityRequest, _entities, _relations, _save_state as _kg_save,
        )

        terms = _extract_key_terms(text)
        if not terms:
            return 0

        linked = 0
        for term in terms:
            # Check if entity already exists
            exists = any(e.get("name", "").lower() == term.lower() for e in _entities.values())
            if exists:
                continue

            eid = str(uuid.uuid4())
            now = _now_iso()
            entity = {
                "id": eid,
                "name": term,
                "type": "concept",
                "aliases": [],
                "description": f"Extracted from document: {title}",
                "properties": {"sourceDocId": doc_id, "sourceDocTitle": title},
                "source": "knowledge_base",
                "confidence": 0.6,
                "createdAt": now,
                "updatedAt": now,
                "accessCount": 0,
                "lastAccessed": now,
            }
            _entities[eid] = entity
            linked += 1

        if linked > 0:
            await _kg_save()
            logger.info("Linked %d terms from doc '%s' to knowledge graph", linked, title)

        return linked
    except Exception as exc:
        logger.debug("Failed to link doc to knowledge graph: %s", exc)
        return 0


# ── Embedding integration helpers ─────────────────────────────────────


async def _compute_embeddings_for_doc(doc_id: str, chunk_texts: list[str]) -> int:
    """Compute embeddings for a document's chunks and persist to NPY file.

    Returns the number of chunks successfully embedded (0 if no provider available).
    """
    global _chunk_vectors_np

    provider = get_embedding_provider()
    if provider is None:
        await init_embedding_provider()
        provider = get_embedding_provider()

    if provider is None or not provider.is_available():
        return 0

    vectors = await provider.embed(chunk_texts)
    if vectors is None:
        logger.warning("Embedding provider returned None for doc %s", doc_id)
        return 0

    if len(vectors) != len(chunk_texts):
        logger.warning("Embedding count mismatch: %d vectors vs %d chunks for doc %s",
                       len(vectors), len(chunk_texts), doc_id)
        return 0

    # Merge new vectors with existing ones (align with chunk_index)
    new_vecs = np.array(vectors, dtype=np.float32)
    if _chunk_vectors_np is not None and _chunk_vectors_np.size > 0:
        _chunk_vectors_np = np.vstack([_chunk_vectors_np, new_vecs])
    else:
        _chunk_vectors_np = new_vecs

    _save_vectors(_chunk_vectors_np)
    logger.info("Embedded %d chunks for doc '%s' (%dd vectors)",
                len(vectors), doc_id, new_vecs.shape[1])
    return len(vectors)


async def _hybrid_search(
    query: str,
    top_k: int = 10,
    doc_id: str | None = None,
    alpha: float = 0.6,
    min_score: float = 0.05,
) -> tuple[list[HybridSearchResult], str, float]:
    """Perform hybrid search combining TF-IDF and embedding similarity.

    hybrid_score = alpha * norm_emb + (1 - alpha) * norm_tfidf

    Returns: (results, search_mode, query_time_ms)
    """
    import time as _time
    t0 = _time.monotonic()

    global _chunk_index, _chunk_vectors_np

    provider = get_embedding_provider()
    has_embeddings = (
        provider is not None
        and provider.is_available()
        and _chunk_vectors_np is not None
        and _chunk_vectors_np.shape[0] == len(_chunk_index)
    )

    # ── Determine search mode ─────────────────────────────────────────
    if alpha >= 0.95 and has_embeddings:
        mode = "embedding_only"
    elif alpha <= 0.05 or not has_embeddings:
        mode = "tfidf_only"
    else:
        mode = "hybrid"

    # ── TF-IDF scores ─────────────────────────────────────────────────
    tfidf_results = _search_chunks(query, top_k=top_k * 5, doc_id=doc_id, min_score=0.0)

    # Build score maps indexed by (doc_id, chunk_idx)
    tfidf_score_map: dict[tuple[str, int], float] = {}
    for r in tfidf_results:
        key = (r.docId, r.chunkIdx)
        tfidf_score_map[key] = r.score

    # ── Embedding scores ──────────────────────────────────────────────
    emb_score_map: dict[tuple[str, int], float] = {}
    if has_embeddings and mode != "tfidf_only":
        query_vecs = await provider.embed([query])
        if query_vecs and query_vecs[0]:
            sims = compute_similarity(query_vecs[0], _chunk_vectors_np)
            for i, sim in enumerate(sims):
                if i < len(_chunk_index):
                    entry = _chunk_index[i]
                    key = (entry["doc_id"], entry["chunk_idx"])
                    emb_score_map[key] = float(max(0, sim))  # clamp negative to 0

    # ── Fuse scores ───────────────────────────────────────────────────
    # Collect candidates with at least one non-zero score
    candidates: dict[tuple[str, int], dict] = {}
    all_keys = set(tfidf_score_map.keys()) | set(emb_score_map.keys())

    for key in all_keys:
        if doc_id and key[0] != doc_id:
            continue
        tfidf_s = tfidf_score_map.get(key, 0.0)
        emb_s = emb_score_map.get(key, 0.0)

        if mode == "tfidf_only":
            hybrid_s = tfidf_s
        elif mode == "embedding_only":
            hybrid_s = emb_s
        else:
            # Min-max normalize each dimension, then weight
            # If only one source contributes, use it directly
            if tfidf_s > 0 and emb_s > 0:
                hybrid_s = alpha * emb_s + (1 - alpha) * tfidf_s
            elif tfidf_s > 0:
                hybrid_s = (1 - alpha) * tfidf_s  # penalize solo TF-IDF
            elif emb_s > 0:
                hybrid_s = alpha * emb_s
            else:
                continue

        if hybrid_s >= min_score:
            candidates[key] = {
                "tfidf_s": round(tfidf_s, 4),
                "emb_s": round(emb_s, 4),
                "hybrid_s": round(hybrid_s, 4),
            }

    # ── Build results ─────────────────────────────────────────────────
    results: list[HybridSearchResult] = []
    sorted_candidates = sorted(candidates.items(), key=lambda x: x[1]["hybrid_s"], reverse=True)

    for key, scores in sorted_candidates[:top_k]:
        doc_id_k, chunk_idx = key
        doc = _documents.get(doc_id_k, {})

        # Find chunk text
        chunk_text = ""
        for i, entry in enumerate(_chunk_index):
            if entry["doc_id"] == doc_id_k and entry["chunk_idx"] == chunk_idx:
                chunk_text = entry["text"]
                break

        preview = chunk_text[:300]
        if len(chunk_text) > 300:
            preview += "..."

        results.append(HybridSearchResult(
            docId=doc_id_k,
            docTitle=doc.get("title", ""),
            docFilename=doc.get("filename", ""),
            chunkIdx=chunk_idx,
            chunkText=chunk_text,
            score=scores["hybrid_s"],
            tfidfScore=scores["tfidf_s"],
            embeddingScore=scores["emb_s"],
            preview=preview,
        ))

    query_time = (_time.monotonic() - t0) * 1000
    return results, mode, query_time


# ── Document endpoints ────────────────────────────────────────────────


@router.post("/documents")
async def upload_document(
    file: UploadFile = File(...),
    title: str | None = None,
    linkToGraph: bool = Query(default=True, alias="linkToGraph"),
    category: str = Query(default="general", max_length=64),
    tags: str = Query(default="", max_length=1024, description="Comma-separated tags"),
) -> DocumentMeta:
    """Upload a document to the knowledge base.

    Extracts text, chunks it, indexes for TF-IDF search, and optionally
    extracts key terms as entities in the Knowledge Graph.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename is required")

    ext = Path(file.filename).suffix.lower()
    if ext not in SUPPORTED_EXTS and ext not in {".pdf", ".docx"}:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Supported: {sorted(SUPPORTED_EXTS | {'.pdf', '.docx'})}"
        )

    # Read file content
    content = await file.read()
    if len(content) > MAX_DOC_SIZE_MB * 1024 * 1024:
        raise HTTPException(status_code=400, detail=f"File exceeds {MAX_DOC_SIZE_MB}MB limit")

    # Save file to docs directory
    doc_id = str(uuid.uuid4())
    docs_dir = _get_docs_dir()
    safe_filename = f"{doc_id}{ext}"
    filepath = docs_dir / safe_filename
    filepath.write_bytes(content)

    # Extract text
    text = _extract_text(str(filepath), ext)
    if text is None:
        filepath.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail="Failed to extract text from file")
    if not text.strip():
        filepath.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail="No text content found in file")

    doc_title = title or Path(file.filename).stem
    now = _now_iso()
    char_count = len(text)

    # Parse tags and infer category
    tag_list = [t.strip().lower() for t in tags.split(",") if t.strip()] if tags else []
    effective_category = category if category != "general" else _infer_category(ext)

    # Get page count for PDF/DOCX
    page_count = _get_page_count(str(filepath), ext)

    # Chunk the text
    chunk_texts = _smart_chunk(text)
    chunks = [
        {"idx": i, "text": ct, "charCount": len(ct), "createdAt": now}
        for i, ct in enumerate(chunk_texts)
    ]

    # If we have PDF/DOCX page count, add page info to chunks
    if page_count and page_count > 0:
        total_pages = page_count
        chars_per_page = char_count / total_pages if total_pages > 0 else 1
        for chunk in chunks:
            # Estimate which page this chunk belongs to
            approx_page = int(chunk["charCount"] / chars_per_page) if chars_per_page > 0 else 0
            chunk["page"] = min(approx_page + 1, total_pages)

    async with _lock:
        # Store document metadata
        _documents[doc_id] = {
            "id": doc_id,
            "filename": file.filename,
            "title": doc_title,
            "fileType": ext,
            "fileSize": len(content),
            "chunkCount": len(chunks),
            "charCount": char_count,
            "pageCount": page_count,
            "tags": tag_list,
            "category": effective_category,
            "createdAt": now,
            "updatedAt": now,
        }
        _chunks[doc_id] = chunks

        # Rebuild TF-IDF index
        _rebuild_index()

        # Persist metadata + chunks
        await _save_state()

    # Compute embeddings for new chunks (outside lock — makes API calls)
    embedding_count = await _compute_embeddings_for_doc(doc_id, chunk_texts)

    if embedding_count > 0:
        # Rebuild index to pick up new vectors
        async with _lock:
            _rebuild_index()

    # Auto-link to knowledge graph (non-blocking background task)
    if linkToGraph:
        asyncio.create_task(_link_to_knowledge_graph(doc_id, doc_title, text))

    logger.info("Ingested document '%s' (%d chars, %d chunks%s) as %s",
                file.filename, char_count, len(chunks),
                f", {embedding_count} embeddings" if embedding_count > 0 else "",
                doc_id)

    return DocumentMeta(
        id=doc_id,
        filename=file.filename,
        title=doc_title,
        fileType=ext,
        fileSize=len(content),
        chunkCount=len(chunks),
        charCount=char_count,
        pageCount=page_count,
        tags=tag_list,
        category=effective_category,
        createdAt=now,
        updatedAt=now,
    )


@router.get("/documents", response_model=DocumentListResponse)
async def list_documents(
    fileType: str | None = Query(default=None),
    category: str | None = Query(default=None),
    tag: str | None = Query(default=None),
    search: str | None = Query(default=None),
    limit: int = Query(default=100, ge=1),
) -> DocumentListResponse:
    """List all documents in the knowledge base with optional filtering."""
    results: list[DocumentMeta] = []
    for d in _documents.values():
        if fileType and d.get("fileType") != fileType:
            continue
        if category and d.get("category", "general") != category:
            continue
        if tag:
            doc_tags = [t.lower() for t in d.get("tags", [])]
            if tag.lower() not in doc_tags:
                continue
        if search:
            hay = f"{d.get('title', '')} {d.get('filename', '')} {' '.join(d.get('tags', []))}".lower()
            if search.lower() not in hay:
                continue
        results.append(DocumentMeta(
            id=d["id"],
            filename=d.get("filename", ""),
            title=d.get("title", ""),
            fileType=d.get("fileType", ""),
            fileSize=d.get("fileSize", 0),
            chunkCount=d.get("chunkCount", 0),
            charCount=d.get("charCount", 0),
            pageCount=d.get("pageCount"),
            tags=d.get("tags", []),
            category=d.get("category", "general"),
            createdAt=d.get("createdAt", ""),
            updatedAt=d.get("updatedAt", ""),
        ))

    return DocumentListResponse(documents=results[:limit], total=len(results))


@router.get("/documents/{doc_id}", response_model=DocumentDetail)
async def get_document(doc_id: str) -> DocumentDetail:
    """Get a document with all its chunks."""
    doc = _documents.get(doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found")

    return DocumentDetail(
        id=doc["id"],
        filename=doc.get("filename", ""),
        title=doc.get("title", ""),
        fileType=doc.get("fileType", ""),
        fileSize=doc.get("fileSize", 0),
        chunkCount=doc.get("chunkCount", 0),
        charCount=doc.get("charCount", 0),
        pageCount=doc.get("pageCount"),
        tags=doc.get("tags", []),
        category=doc.get("category", "general"),
        createdAt=doc.get("createdAt", ""),
        updatedAt=doc.get("updatedAt", ""),
        chunks=_chunks.get(doc_id, []),
    )


@router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str) -> dict:
    """Delete a document and all its chunks."""
    async with _lock:
        if doc_id not in _documents:
            raise HTTPException(status_code=404, detail="Document not found")

        del _documents[doc_id]
        _chunks.pop(doc_id, None)

        # Remove source file
        docs_dir = _get_docs_dir()
        for ext in SUPPORTED_EXTS | {".pdf", ".docx"}:
            fp = docs_dir / f"{doc_id}{ext}"
            if fp.exists():
                fp.unlink(missing_ok=True)
                break

        # Rebuild index
        _rebuild_index()

        # Persist
        await _save_state()

    return {"success": True}


@router.get("/documents/{doc_id}/download")
async def download_document(doc_id: str):
    """Download the original uploaded document file.

    Returns the file with its original filename via Content-Disposition.
    """
    doc = _documents.get(doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found")

    docs_dir = _get_docs_dir()
    ext = doc.get("fileType", "")
    original_filename = doc.get("filename", f"{doc_id}{ext}")

    # Try looking for the file with the stored extension first
    filepath = docs_dir / f"{doc_id}{ext}"
    if not filepath.exists():
        # Fallback: scan all supported extensions
        for ex in SUPPORTED_EXTS | {".pdf", ".docx"}:
            fp = docs_dir / f"{doc_id}{ex}"
            if fp.exists():
                filepath = fp
                break

    if not filepath.exists():
        raise HTTPException(
            status_code=404,
            detail="Original file not found on disk — it may have been cleaned up",
        )

    # Determine MIME type
    media_type_map = {
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".py": "text/x-python",
        ".js": "application/javascript",
        ".ts": "text/typescript",
        ".json": "application/json",
        ".html": "text/html",
        ".css": "text/css",
        ".yaml": "text/yaml",
        ".yml": "text/yaml",
        ".xml": "application/xml",
        ".csv": "text/csv",
        ".log": "text/plain",
        ".rst": "text/x-rst",
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    media_type = media_type_map.get(ext, "application/octet-stream")

    return FileResponse(
        path=str(filepath),
        filename=original_filename,
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{original_filename}"'},
    )


@router.get("/documents/{doc_id}/related-entities")
async def get_document_related_entities(doc_id: str) -> list[dict]:
    """Return all KG entities that were extracted from this document.

    Delegates to the knowledge_graph module's in-memory store
    by filtering entities whose properties.sourceDocId matches doc_id.
    """
    from app.gateway.routers.knowledge_graph import _entities as kg_entities

    doc = _documents.get(doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found")

    results: list[dict] = []
    for e in kg_entities.values():
        props = e.get("properties", {})
        if isinstance(props, dict) and props.get("sourceDocId") == doc_id:
            results.append(dict(e))
    return results


@router.post("/documents/batch-delete", response_model=BatchDeleteResponse)
async def batch_delete_documents(req: BatchDeleteRequest) -> BatchDeleteResponse:
    """Delete multiple documents at once."""
    deleted = 0
    failed = 0
    errors: list[str] = []

    async with _lock:
        for doc_id in req.ids:
            if doc_id not in _documents:
                failed += 1
                errors.append(f"Document not found: {doc_id}")
                continue

            del _documents[doc_id]
            _chunks.pop(doc_id, None)

            # Remove source file
            docs_dir = _get_docs_dir()
            for ext in SUPPORTED_EXTS | {".pdf", ".docx"}:
                fp = docs_dir / f"{doc_id}{ext}"
                if fp.exists():
                    fp.unlink(missing_ok=True)
                    break

            deleted += 1

        # Rebuild index once after all deletes
        _rebuild_index()

        # Persist once
        await _save_state()

    return BatchDeleteResponse(
        success=len(errors) == 0,
        deleted=deleted,
        failed=failed,
        errors=errors,
    )


@router.post("/documents/batch-update", response_model=BatchUpdateResponse)
async def batch_update_documents(req: BatchUpdateRequest) -> BatchUpdateResponse:
    """Update tags and/or category for multiple documents at once."""
    updated = 0
    failed = 0
    errors: list[str] = []

    valid_modes = {"set", "add", "remove"}
    mode = req.mode if req.mode in valid_modes else "set"

    async with _lock:
        for doc_id in req.ids:
            doc = _documents.get(doc_id)
            if doc is None:
                failed += 1
                errors.append(f"Document not found: {doc_id}")
                continue

            # Apply tags based on mode
            if req.tags is not None:
                clean_tags = [t.strip().lower() for t in req.tags if t.strip()]
                if mode == "set":
                    doc["tags"] = clean_tags
                elif mode == "add":
                    existing = doc.get("tags", [])
                    for t in clean_tags:
                        if t not in existing:
                            existing.append(t)
                    doc["tags"] = existing
                elif mode == "remove":
                    existing = doc.get("tags", [])
                    doc["tags"] = [t for t in existing if t not in clean_tags]

            # Apply category
            if req.category is not None:
                doc["category"] = req.category.strip().lower()

            doc["updatedAt"] = _now_iso()
            updated += 1

        if updated > 0:
            await _save_state()

    return BatchUpdateResponse(
        success=len(errors) == 0,
        updated=updated,
        failed=failed,
        errors=errors,
    )


@router.patch("/documents/{doc_id}", response_model=DocumentMeta)
async def update_document(doc_id: str, req: DocumentUpdateRequest) -> DocumentMeta:
    """Update document metadata (title, tags, category)."""
    async with _lock:
        doc = _documents.get(doc_id)
        if doc is None:
            raise HTTPException(status_code=404, detail="Document not found")

        if req.title is not None:
            doc["title"] = req.title.strip()
        if req.tags is not None:
            doc["tags"] = [t.strip().lower() for t in req.tags if t.strip()]
        if req.category is not None:
            doc["category"] = req.category.strip().lower()
        doc["updatedAt"] = _now_iso()

        await _save_state()

    return DocumentMeta(
        id=doc["id"],
        filename=doc.get("filename", ""),
        title=doc.get("title", ""),
        fileType=doc.get("fileType", ""),
        fileSize=doc.get("fileSize", 0),
        chunkCount=doc.get("chunkCount", 0),
        charCount=doc.get("charCount", 0),
        pageCount=doc.get("pageCount"),
        tags=doc.get("tags", []),
        category=doc.get("category", "general"),
        createdAt=doc.get("createdAt", ""),
        updatedAt=doc.get("updatedAt", ""),
    )


# ── Tags / Categories ──────────────────────────────────────────────


def _infer_category(ext: str) -> str:
    """Infer a category from file extension."""
    cat_map = {
        ".py": "code", ".js": "code", ".ts": "code", ".css": "code",
        ".html": "code", ".json": "code", ".xml": "code",
        ".md": "documentation", ".rst": "documentation",
        ".txt": "general", ".csv": "data", ".log": "data",
        ".yaml": "code", ".yml": "code",
        ".pdf": "research", ".docx": "documentation",
    }
    return cat_map.get(ext, "general")


@router.get("/tags", response_model=TagsResponse)
async def list_tags() -> TagsResponse:
    """List all tags and categories with document counts."""
    tag_counts: dict[str, int] = {}
    cat_counts: dict[str, int] = {}

    for d in _documents.values():
        cat = d.get("category", "general")
        cat_counts[cat] = cat_counts.get(cat, 0) + 1
        for t in d.get("tags", []):
            if t:
                tag_counts[t] = tag_counts.get(t, 0) + 1

    tags_list = sorted(
        [{"name": k, "count": v} for k, v in tag_counts.items()],
        key=lambda x: x["count"], reverse=True,
    )
    cats_list = sorted(
        [{"name": k, "count": v} for k, v in cat_counts.items()],
        key=lambda x: x["count"], reverse=True,
    )

    return TagsResponse(tags=tags_list, categories=cats_list)


# ── Search endpoint ───────────────────────────────────────────────────


@router.post("/search", response_model=SearchResponse)
async def search_knowledge_base(req: SearchRequest) -> SearchResponse:
    """Semantic search across the knowledge base using TF-IDF + cosine similarity."""
    start = time.time()

    results = _search_chunks(
        query=req.query,
        top_k=req.topK,
        doc_id=req.docId,
        min_score=req.minScore,
    )

    query_time = (time.time() - start) * 1000

    return SearchResponse(
        results=results,
        totalChunks=len(_chunk_index),
        queryTimeMs=round(query_time, 2),
    )


# ── Hybrid search endpoint ───────────────────────────────────────────


@router.post("/search/hybrid", response_model=HybridSearchResponse)
async def hybrid_search_knowledge_base(req: HybridSearchRequest) -> HybridSearchResponse:
    """Hybrid semantic search: weighted fusion of TF-IDF + Embedding similarity.

    When an embedding provider is configured (e.g. OpenAI text-embedding-3-small),
    results combine both lexical (TF-IDF) and semantic (embedding) relevance.
    The alpha parameter controls the balance:

        alpha = 0.0  →  TF-IDF only
        alpha = 0.6  →  balanced hybrid (default)
        alpha = 1.0  →  embedding only

    Falls back to TF-IDF-only when no embedding provider is available.
    """
    results, mode, query_time = await _hybrid_search(
        query=req.query,
        top_k=req.topK,
        doc_id=req.docId,
        alpha=req.alpha,
        min_score=req.minScore,
    )

    return HybridSearchResponse(
        results=results,
        totalChunks=len(_chunk_index),
        queryTimeMs=round(query_time, 2),
        searchMode=mode,
        alpha=req.alpha,
    )


# ── Embedding status endpoint ────────────────────────────────────────


@router.get("/embeddings/status", response_model=EmbeddingStatusResponse)
async def get_embedding_status() -> EmbeddingStatusResponse:
    """Return current embedding provider status and vector stats."""
    from app.gateway.routers.embeddings import get_embedding_status as _get_status
    status = await _get_status()

    return EmbeddingStatusResponse(
        provider=status.get("provider", "none"),
        model=status.get("model", "N/A"),
        dimension=status.get("dimension", 0),
        available=status.get("available", False),
        embeddedChunks=_chunk_vectors_np.shape[0] if _chunk_vectors_np is not None else 0,
        message=status.get("message", ""),
    )


# ── Stats endpoint ────────────────────────────────────────────────────


@router.get("/stats", response_model=KBStatsResponse)
async def get_stats() -> KBStatsResponse:
    """Get knowledge base statistics."""
    file_types: dict[str, int] = {}
    categories: dict[str, int] = {}
    tag_counts: dict[str, int] = {}
    total_chars = 0
    total_chunks_count = 0

    for d in _documents.values():
        ft = d.get("fileType", "unknown")
        file_types[ft] = file_types.get(ft, 0) + 1
        cat = d.get("category", "general")
        categories[cat] = categories.get(cat, 0) + 1
        for t in d.get("tags", []):
            if t:
                tag_counts[t] = tag_counts.get(t, 0) + 1
        total_chars += d.get("charCount", 0)
        total_chunks_count += d.get("chunkCount", 0)

    avg_chunk = round(total_chars / total_chunks_count, 1) if total_chunks_count > 0 else 0.0
    indexed = len(_chunk_index)

    return KBStatsResponse(
        totalDocuments=len(_documents),
        totalChunks=total_chunks_count,
        totalChars=total_chars,
        fileTypes=file_types,
        indexedVectors=indexed,
        avgChunkSize=avg_chunk,
        categories=categories,
        tags=tag_counts,
    )


# ── Reindex endpoints ───────────────────────────────────────────────


@router.post("/documents/{doc_id}/reindex", response_model=DocumentMeta)
async def reindex_document(doc_id: str) -> DocumentMeta:
    """Re-extract text, re-chunk, and re-index a single document."""
    async with _lock:
        doc = _documents.get(doc_id)
        if doc is None:
            raise HTTPException(status_code=404, detail="Document not found")

        docs_dir = _get_docs_dir()
        src_path: str | None = None
        src_ext: str | None = None

        # Find source file on disk
        for ext in SUPPORTED_EXTS | {".pdf", ".docx"}:
            fp = docs_dir / f"{doc_id}{ext}"
            if fp.exists():
                src_path = str(fp)
                src_ext = ext
                break

        if src_path is None or src_ext is None:
            raise HTTPException(status_code=404, detail="Source file not found on disk")

        # Re-extract text
        text = _extract_text(src_path, src_ext)
        if text is None:
            raise HTTPException(status_code=500, detail="Failed to extract text from source file")

        # Re-chunk
        new_chunks = _smart_chunk(text)
        if not new_chunks:
            raise HTTPException(status_code=500, detail="No text content could be extracted")

        # Update chunks
        chunk_list = []
        total_chars = 0
        now_iso = _now_iso()
        for i, chunk_text in enumerate(new_chunks):
            chunk_list.append({
                "idx": i,
                "text": chunk_text,
                "charCount": len(chunk_text),
                "createdAt": now_iso,
            })
            total_chars += len(chunk_text)

        _chunks[doc_id] = chunk_list

        # Update document metadata
        doc["chunkCount"] = len(new_chunks)
        doc["charCount"] = total_chars
        doc["updatedAt"] = now_iso

        # Rebuild index and persist
        _rebuild_index()
        await _save_state()

    # Recompute embeddings in background
    if embedding_available():
        asyncio.create_task(_compute_embeddings_for_doc(doc_id))

    return DocumentMeta(
        id=doc["id"],
        filename=doc.get("filename", ""),
        title=doc.get("title", ""),
        fileType=doc.get("fileType", ""),
        fileSize=doc.get("fileSize", 0),
        chunkCount=doc.get("chunkCount", 0),
        charCount=doc.get("charCount", 0),
        pageCount=doc.get("pageCount"),
        tags=doc.get("tags", []),
        category=doc.get("category", "general"),
        createdAt=doc.get("createdAt", ""),
        updatedAt=doc.get("updatedAt", ""),
    )


@router.post("/reindex")
async def reindex_knowledge_base() -> dict:
    """Force rebuild the TF-IDF index from all stored chunks."""
    async with _lock:
        _rebuild_index()
    return {
        "success": True,
        "indexedVectors": len(_chunk_index),
        "message": f"Reindexed {len(_chunk_index)} chunks",
    }
